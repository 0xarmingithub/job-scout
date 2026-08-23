"""
cv_import.py — draft a profile.yaml from a CV.

Writing profile.yaml by hand takes half an hour. Most of what goes in it is
already written down in your CV, so this reads the CV and drafts the file for
you.

    job-scout init ~/job-search --from-cv ~/cv.pdf

Know what this sends where. The text of your CV is sent to whichever model
`scoring_model` names, which for the default is Google's API. That text has your
name, your address, your phone number and your whole employment history in it.
If you would rather it stayed on your machine, use a CLI backend you run locally
(`claude:`, `grok:` or `codex:`) or write profile.yaml by hand. The scout never
stores or transmits the CV anywhere else, and does not keep a copy.

One thing it deliberately will not do: fill in `confirmed_gaps`. A CV lists what
you have done. It cannot tell anyone what you cannot do, and guessing would be
worse than leaving it empty, because a wrong gap silently caps good jobs at 40.
That section is left for you, and the command says so when it finishes.

Formats: .txt and .md need nothing. .pdf needs pypdf. .docx needs python-docx.
Both are in the `cv` extra:

    pip install -e ".[cv]"

If you would rather not install either, open the CV, select all, and paste it
into a .txt file.
"""

import logging
from pathlib import Path

import yaml

from .llm import ModelError, label_for, preflight, run_model

logger = logging.getLogger(__name__)

# Enough for a long CV. Past this the model is reading the same career twice.
MAX_CV_CHARS = 20_000

SYSTEM_PROMPT = (
    "You convert a CV into a YAML profile. You output only YAML, with no "
    "commentary and no markdown fences. You never invent experience."
)


class CvImportError(RuntimeError):
    """The CV could not be read or the model could not draft a profile."""


# ─── Reading the file ─────────────────────────────────────────────────────────

def read_cv_text(path: Path) -> str:
    """Return the plain text of a CV. Raises CvImportError with a fix in it."""
    path = Path(path).expanduser()
    if not path.exists():
        raise CvImportError(f"No CV at {path}.")

    suffix = path.suffix.lower()
    if suffix in (".txt", ".md", ""):
        text = path.read_text(encoding="utf-8", errors="replace")
    elif suffix == ".pdf":
        text = _read_pdf(path)
    elif suffix == ".docx":
        text = _read_docx(path)
    else:
        raise CvImportError(
            f"Cannot read {suffix or 'that'} files. Supported: .txt, .md, .pdf, "
            f".docx. Open the CV, select all, and paste it into a .txt file."
        )

    text = text.strip()
    if len(text) < 200:
        raise CvImportError(
            f"{path} gave only {len(text)} characters of text. If it is a "
            f"scanned or image-only PDF there is nothing to extract. Paste the "
            f"text into a .txt file instead."
        )
    if len(text) > MAX_CV_CHARS:
        logger.warning(
            "%s is %d characters; only the first %d were used. If your most "
            "recent roles are at the end of the file, reorder it or trim the "
            "older ones.",
            path, len(text), MAX_CV_CHARS,
        )
    return text[:MAX_CV_CHARS]


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise CvImportError(
            "Reading a PDF needs the pypdf package, which is not installed. "
            'Install it with: pip install pypdf  (or pip install -e ".[cv]")'
        ) from exc

    try:
        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:
        raise CvImportError(f"Could not read {path}: {exc}") from exc


def _read_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise CvImportError(
            "Reading a .docx needs the python-docx package, which is not "
            'installed. Install it with: pip install python-docx  (or pip '
            'install -e ".[cv]")'
        ) from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise CvImportError(f"Could not read {path}: {exc}") from exc

    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


# ─── Asking the model ─────────────────────────────────────────────────────────

def build_prompt(cv_text: str) -> str:
    return (
        "Below is somebody's CV. Turn it into a YAML profile for a job-matching "
        "tool, using the exact keys shown in the template.\n\n"
        "RULES\n"
        "1. Use only what the CV actually says. Never add a skill, employer, "
        "date or level that is not in the text.\n"
        "2. Leave confirmed_gaps as an empty list. A CV says what somebody has "
        "done; it cannot say what they cannot do, and a wrong entry there is "
        "worse than none.\n"
        "3. core_skills are things the CV shows real, repeated, recent use of. "
        "secondary_skills are things it mentions once or in passing.\n"
        "4. target_roles: list every job title this person's experience fits, "
        "including the different names the same job goes by at different "
        "companies.\n"
        "5. extra_pre_filter_keywords: 30 to 50 single words or short phrases "
        "drawn from the CV's technologies and domains. Lowercase.\n"
        "6. If the CV does not say something, write the string "
        "\"not stated\" rather than guessing. Leave "
        "hard_exclude_location_patterns and hard_exclude_title_patterns as "
        "empty lists.\n"
        "7. Output YAML only. No fences, no commentary.\n\n"
        "TEMPLATE\n"
        "candidate:\n"
        "  name:\n"
        "  current_role:\n"
        "  years_experience:\n"
        "  seniority:\n"
        "  location:\n"
        "  work_authorization:\n"
        "  target_geography:\n"
        "  languages:\n"
        "    English:\n"
        "target_roles: []\n"
        "core_skills: []\n"
        "secondary_skills: []\n"
        "confirmed_gaps: []\n"
        "industries_preferred: []\n"
        "extra_pre_filter_keywords: []\n"
        "hard_exclude_location_patterns: []\n"
        "hard_exclude_title_patterns: []\n\n"
        "CV\n"
        f"{cv_text}"
    )


def _strip_fences(text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        lines = cleaned.splitlines()
        lines = lines[1:]
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        cleaned = "\n".join(lines)
    return cleaned.strip()


def profile_from_cv(cv_path: Path, model_spec: str) -> tuple[str, dict]:
    """
    Read the CV, ask the model for a profile, and return (yaml_text, parsed).

    Raises CvImportError with something actionable in the message.
    """
    problem = preflight(model_spec)
    if problem:
        raise CvImportError(
            f"Reading a CV needs a working model backend, and {problem}"
        )

    cv_text = read_cv_text(cv_path)
    logger.info("Read %d characters from %s", len(cv_text), cv_path)
    logger.info("Drafting a profile with %s", label_for(model_spec))

    try:
        reply = run_model(
            model_spec, SYSTEM_PROMPT, build_prompt(cv_text), max_tokens=4096
        )
    except ModelError as exc:
        raise CvImportError(f"The model call failed: {exc}") from exc

    text = _strip_fences(reply)
    try:
        parsed = yaml.safe_load(text)
    except yaml.YAMLError as exc:
        raise CvImportError(
            f"The model did not return valid YAML: {exc}\n"
            f"Try again, or use a stronger model in scoring_model."
        ) from exc

    if not isinstance(parsed, dict) or not parsed.get("candidate"):
        raise CvImportError(
            "The model's reply had no `candidate` block. Try again, or use a "
            "stronger model in scoring_model."
        )

    # The model is told to leave gaps empty. If it invented some anyway, drop
    # them: a wrong gap silently caps good jobs at 40 and is hard to notice.
    if parsed.get("confirmed_gaps"):
        logger.warning(
            "The model filled in confirmed_gaps from the CV. Dropping them — a "
            "CV cannot show what somebody cannot do."
        )
        parsed["confirmed_gaps"] = []

    return yaml.safe_dump(parsed, sort_keys=False, allow_unicode=True, width=88), parsed


# ─── What to tell the user afterwards ─────────────────────────────────────────

def review_notes(parsed: dict, profile_path: Path) -> str:
    """The message printed after a draft is written. Blunt on purpose."""
    candidate = parsed.get("candidate") or {}
    counts = (
        f"{len(parsed.get('core_skills') or [])} core skills, "
        f"{len(parsed.get('secondary_skills') or [])} secondary, "
        f"{len(parsed.get('target_roles') or [])} target roles, "
        f"{len(parsed.get('extra_pre_filter_keywords') or [])} keywords"
    )
    return (
        f"Drafted {profile_path} for {candidate.get('name', 'the candidate')}.\n"
        f"  {counts}\n"
        "\n"
        "This is a draft from your CV, and three things in it are wrong or\n"
        "missing by construction. Fix them before you trust a score.\n"
        "\n"
        "1. confirmed_gaps is empty, and it is the most valuable section in the\n"
        "   file. Your CV cannot say what you cannot do. Write down the things\n"
        "   that come up in adverts in your field that you have genuinely never\n"
        "   done. Without it the scorer will send you machine-learning roles\n"
        "   because it saw the word Python.\n"
        "\n"
        "2. hard_exclude_location_patterns is empty. Add the places you will\n"
        "   not commute to.\n"
        "\n"
        "3. Read candidate.work_authorization and candidate.languages. A CV\n"
        "   rarely states either properly, and they are what reject a job\n"
        "   outright rather than scoring it low.\n"
    )
